"""
domain/ingestion.py — Pure parsers + text splitting. No Qdrant, no storage, no registry.
"""

import functools
import html
import logging
import re
from pathlib import Path

from config import settings
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter

log = logging.getLogger("detailed")


def setup_logger() -> logging.Logger:
    logger = logging.getLogger("detailed")
    if logger.handlers:
        return logger

    logger.setLevel(logging.DEBUG)

    fmt = logging.Formatter(
        fmt="%(asctime)s  %(levelname)-8s  %(message)s",
        datefmt="%H:%M:%S",
    )

    ch = logging.StreamHandler()
    ch.setLevel(logging.DEBUG)
    ch.setFormatter(fmt)
    logger.addHandler(ch)

    log_path = Path(settings.data_dir) / "ingestion.log"
    try:
        log_path.parent.mkdir(parents=True, exist_ok=True)
        fh = logging.FileHandler(log_path, encoding="utf-8")
        fh.setLevel(logging.DEBUG)
        fh.setFormatter(fmt)
        logger.addHandler(fh)
    except OSError:
        pass

    return logger


# ---------------------------------------------------------------------------
# OCR — lazy-loaded via lru_cache (no global keyword)
# ---------------------------------------------------------------------------


@functools.lru_cache(maxsize=1)
def _get_paddle_ocr():
    from paddleocr import PaddleOCR

    log.info("Loading PaddleOCR (lang=%s) ...", settings.ocr_lang_paddle)
    return PaddleOCR(
        use_angle_cls=True,
        lang=settings.ocr_lang_paddle,
        show_log=False,
    )


@functools.lru_cache(maxsize=1)
def _get_surya_predictors():
    from surya.detection import DetectionPredictor
    from surya.recognition import RecognitionPredictor

    log.info("Loading Surya OCR ...")
    return (RecognitionPredictor(), DetectionPredictor())


def _ocr_image_paddle(image) -> str:
    import numpy as np

    ocr = _get_paddle_ocr()
    result = ocr.ocr(np.array(image), cls=True)
    lines = []
    for block in result or []:
        for entry in block or []:
            text = entry[1][0]
            if text and text.strip():
                lines.append(text.strip())
    return "\n".join(lines)


def _ocr_image_surya(image) -> str:
    rec_predictor, det_predictor = _get_surya_predictors()
    predictions = rec_predictor([image], [settings.ocr_lang_surya], det_predictor)
    lines = [line.text.strip() for line in predictions[0].text_lines if line.text.strip()]
    return "\n".join(lines)


def ocr_pdf_pages(doc, page_nums: list[int], filename: str) -> dict:
    import fitz
    from PIL import Image

    results = {}
    zoom = settings.ocr_dpi / 72
    mat = fitz.Matrix(zoom, zoom)

    for page_num in page_nums:
        page = doc.load_page(page_num - 1)
        pix = page.get_pixmap(matrix=mat)
        img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)

        text = ""
        if settings.ocr_engine in ("paddleocr", "auto"):
            text = _ocr_image_paddle(img)

        if not text and settings.ocr_engine in ("surya", "auto"):
            text = _ocr_image_surya(img)

        results[page_num] = text

    return results


# ---------------------------------------------------------------------------
# PDF parsing
# ---------------------------------------------------------------------------


def parse_pdf(file_path: Path) -> list[Document]:
    import fitz

    doc = fitz.open(str(file_path))
    pages = []

    ocr_pages_needed = []
    for page_num in range(1, len(doc) + 1):
        page = doc.load_page(page_num - 1)
        text = page.get_text("text").strip()
        if not text and settings.ocr_enabled:
            ocr_pages_needed.append(page_num)
        elif text:
            pages.append(
                Document(
                    page_content=text,
                    metadata={"page": page_num, "source": str(file_path)},
                )
            )

    if ocr_pages_needed:
        ocr_results = ocr_pdf_pages(doc, ocr_pages_needed, file_path.name)
        for page_num, text in ocr_results.items():
            if text:
                pages.append(
                    Document(
                        page_content=text,
                        metadata={"page": page_num, "source": str(file_path)},
                    )
                )

    doc.close()
    return pages


# ---------------------------------------------------------------------------
# Text splitting
# ---------------------------------------------------------------------------


def split_documents(docs: list[Document]) -> list[Document]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.chunk_size,
        chunk_overlap=settings.chunk_overlap,
        length_function=len,
    )
    chunks = splitter.split_documents(docs)
    log.info("Split %d documents into %d chunks", len(docs), len(chunks))
    return chunks


# ---------------------------------------------------------------------------
# Parsers for non-PDF formats
# ---------------------------------------------------------------------------


def _parse_docx(file_path: Path) -> str:
    import docx

    doc = docx.Document(str(file_path))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def _parse_rtf(file_path: Path) -> str:
    from striprtf.striprtf import rtf_to_text

    return rtf_to_text(file_path.read_text(encoding="utf-8", errors="replace"))


def _parse_markdown(file_path: Path) -> str:
    text = file_path.read_text(encoding="utf-8", errors="replace")
    text = re.sub(r"!\[.*?\]\(.*?\)", "", text)
    text = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", text)
    text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)
    text = re.sub(r"[*_`~]+", "", text)
    text = html.unescape(text)
    return text.strip()


def _parse_txt(file_path: Path) -> str:
    return file_path.read_text(encoding="utf-8", errors="replace")


PARSERS = {
    ".docx": _parse_docx,
    ".doc": _parse_docx,
    ".rtf": _parse_rtf,
    ".md": _parse_markdown,
    ".txt": _parse_txt,
}


# ---------------------------------------------------------------------------
# Text cleaning
# ---------------------------------------------------------------------------


def clean_pdf_text(text: str) -> str:
    """Clean extracted PDF text: fix hyphenation, collapse whitespace, remove decorative lines."""
    # Fix hyphenation at line breaks
    text = re.sub(r"-\n", "", text)
    # Collapse multiple whitespace to single space
    text = re.sub(r"[^\S\n]+", " ", text)
    # Remove decorative separator lines (---, ===, *** etc.)
    text = re.sub(r"^[•\-=~*]{3,}\s*$", "", text, flags=re.MULTILINE)
    # Collapse multiple blank lines
    text = re.sub(r"\n{2,}", "\n", text)
    return text.strip()
