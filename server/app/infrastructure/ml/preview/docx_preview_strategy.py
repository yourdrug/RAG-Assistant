"""DOCX preview strategy — native parsing via python-docx + mammoth for images."""

from __future__ import annotations

import io
import logging
from pathlib import Path

from domain.value_objects.page_content_type import PageContentType, PreviewUnitKind

from application.services.pdf_diagnostic_service import DryRunPageResult
from infrastructure.ml.pdf_diag import is_garbled
from infrastructure.ml.ingestion import parse_docx_sections

logger = logging.getLogger("default")

_DRY_RUN_EXTENSIONS = {".docx"}


def _extract_images_without_alt(docx_path: Path) -> list[bytes]:
    """Extract raw bytes of images that have no alt-text via mammoth."""
    import mammoth

    images: list[bytes] = []

    def _handler(image: object) -> dict:
        alt = getattr(image, "alt_text", None)
        if not alt:
            opener = getattr(image, "open", None)
            if opener is not None:
                with opener() as f:
                    images.append(f.read())
        return {}

    with open(docx_path, "rb") as f:
        mammoth.convert_to_html(f, convert_image=_handler)
    return images


def _has_images_without_text(docx_path: Path) -> list[bytes]:
    """Check for paragraphs with drawings but no text content."""
    from docx.oxml.ns import qn

    doc = DocxDocumentWrapper(docx_path)
    image_blobs: list[bytes] = []

    for p in doc.paragraphs:
        has_text = bool(p.text.strip())
        has_drawing = bool(p._element.findall(f".//{qn('w:drawing')}"))
        if has_drawing and not has_text:
            blob = _extract_first_image_from_paragraph(p, doc)
            if blob is not None:
                image_blobs.append(blob)

    return image_blobs


class DocxDocumentWrapper:
    """Thin wrapper around python-docx Document for reuse."""

    def __init__(self, path: Path) -> None:
        import docx

        self._doc = docx.Document(str(path))
        self._path = path

    @property
    def paragraphs(self):
        return self._doc.paragraphs

    @property
    def tables(self):
        return self._doc.tables

    @property
    def part(self):
        return self._doc.part


def _extract_first_image_from_paragraph(paragraph, doc_wrapper: DocxDocumentWrapper) -> bytes | None:
    """Extract image bytes from the first <w:drawing> in a paragraph."""
    from docx.oxml.ns import qn

    for drawing in paragraph._element.findall(f".//{qn('w:drawing')}"):
        for blip in drawing.findall(f".//{qn('a:blip')}"):
            r_id = blip.get(qn("r:embed"))
            if r_id and r_id in doc_wrapper.part.rels:
                rel = doc_wrapper.part.rels[r_id]
                try:
                    return rel.target_part.blob
                except Exception:
                    continue
    return None


class DocxPreviewStrategy:
    def supports(self, extension: str) -> bool:
        return extension in _DRY_RUN_EXTENSIONS

    def analyze(self, path: Path) -> tuple[list[DryRunPageResult], dict[str, int], int]:
        sections = parse_docx_sections(path)
        if not sections:
            sections = [(None, "")]

        image_only_blobs = _extract_images_without_alt(path)

        units: list[DryRunPageResult] = []
        types_count: dict[str, int] = {
            PageContentType.TEXT: 0,
            PageContentType.SCAN: 0,
            PageContentType.GARBLED: 0,
            PageContentType.EMPTY: 0,
            PageContentType.TABLE: 0,
            PageContentType.IMAGE_ONLY: 0,
        }
        total_chars = 0

        img_idx = 0
        for idx, (heading, content) in enumerate(sections):
            stripped = content.strip()
            chars = len(stripped)
            total_chars += chars

            is_table = stripped.startswith("\x00TABLE:")
            if is_table:
                stripped = stripped[len("\x00TABLE:") :].strip()
                chars = len(stripped)
                ptype = PageContentType.TABLE
                label = f"Таблица {idx + 1}"
            elif chars == 0:
                if img_idx < len(image_only_blobs):
                    ptype = PageContentType.IMAGE_ONLY
                    label = f"Изображение {img_idx + 1}"
                    img_idx += 1
                else:
                    ptype = PageContentType.EMPTY
                    label = heading or f"Блок {idx + 1}"
            elif is_garbled(stripped):
                ptype = PageContentType.GARBLED
                label = heading or f"Блок {idx + 1}"
            else:
                ptype = PageContentType.TEXT
                label = f"Раздел: {heading}" if heading else f"Блок {idx + 1}"

            types_count[ptype] = types_count.get(ptype, 0) + 1

            preview = stripped[:200] if stripped else ""
            units.append(
                DryRunPageResult(
                    page=idx + 1,
                    type=ptype,
                    content_type=PageContentType.TABLE if is_table else PageContentType.TEXT,
                    chars=chars,
                    preview=preview,
                    full_text=stripped,
                    unit_kind=PreviewUnitKind.SECTION,
                    label=label,
                )
            )

        if not units:
            units.append(
                DryRunPageResult(
                    page=1,
                    type=PageContentType.EMPTY,
                    chars=0,
                    unit_kind=PreviewUnitKind.SECTION,
                    label="Документ пуст",
                )
            )
            types_count[PageContentType.EMPTY] = 1

        return units, types_count, total_chars

    def ocr_problem_units(
        self,
        path: Path,
        units: list[DryRunPageResult],
        unit_ids: list[int],
    ) -> tuple[list[DryRunPageResult], dict[str, int], int]:
        from infrastructure.ml.ingestion import _ocr_image_paddle, _ocr_image_surya
        from config import settings
        from PIL import Image

        image_only_blobs = _extract_images_without_alt(path)

        problem_units = [u for u in units if u.page in unit_ids and u.type == PageContentType.IMAGE_ONLY]
        if not problem_units:
            # Compute summary from existing units even when nothing to OCR
            types_count: dict[str, int] = {
                PageContentType.TEXT: 0,
                PageContentType.SCAN: 0,
                PageContentType.GARBLED: 0,
                PageContentType.EMPTY: 0,
                PageContentType.TABLE: 0,
                PageContentType.IMAGE_ONLY: 0,
            }
            total_chars = 0
            for u in units:
                types_count[u.type] = types_count.get(u.type, 0) + 1
                total_chars += u.chars
            return units, types_count, total_chars

        ocr_results: dict[int, str] = {}
        for u in problem_units:
            blob_idx = u.page - 1
            if blob_idx < len(image_only_blobs):
                blob = image_only_blobs[blob_idx]
                try:
                    img = Image.open(io.BytesIO(blob))
                    if settings.ocr_engine in ("paddleocr", "auto"):
                        text = _ocr_image_paddle(img)
                    else:
                        text = _ocr_image_surya(img)
                    ocr_results[u.page] = text
                except Exception:
                    logger.warning("OCR failed for image unit %d", u.page)

        new_types: dict[str, int] = {
            PageContentType.TEXT: 0,
            PageContentType.SCAN: 0,
            PageContentType.GARBLED: 0,
            PageContentType.EMPTY: 0,
            PageContentType.TABLE: 0,
            PageContentType.IMAGE_ONLY: 0,
        }
        total_chars = 0

        merged: list[DryRunPageResult] = []
        for u in units:
            if u.page in ocr_results and ocr_results[u.page]:
                ocr_text = ocr_results[u.page].strip()
                if ocr_text:
                    merged.append(
                        DryRunPageResult(
                            page=u.page,
                            type=PageContentType.TEXT,
                            content_type=PageContentType.OCR,
                            chars=len(ocr_text),
                            preview=ocr_text[:200],
                            full_text=ocr_text,
                            previous_type=u.type,
                            unit_kind=u.unit_kind,
                            label=u.label,
                        )
                    )
                else:
                    merged.append(u)
            else:
                merged.append(u)
            new_types[merged[-1].type] = new_types.get(merged[-1].type, 0) + 1
            total_chars += merged[-1].chars

        return merged, new_types, total_chars
