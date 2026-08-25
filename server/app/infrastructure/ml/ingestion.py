"""Pure document parsers and text splitting -- no Qdrant, no storage, no registry.

Provides file-type-specific parsers (PDF via PyMuPDF, DOCX, Markdown, plain
text) and a LangChain-compatible text splitter.  All functions are pure:
they accept a file path and return ``langchain.schema.Document`` lists with
no side effects.
"""

import functools
import html
import logging
import re
from pathlib import Path

import docx
import fitz
import numpy as np
from config import settings
from domain.value_objects.doc_domain import DocDomain
from domain.value_objects.page_content_type import PageContentType
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from PIL import Image
from striprtf.striprtf import rtf_to_text

log = logging.getLogger("detailed")


def setup_logger() -> logging.Logger:
    logger = logging.getLogger("detailed")
    if logger.handlers:
        return logger

    logger.setLevel(logging.DEBUG)

    fmt = logging.Formatter(
        fmt="%(asctime)s  %(levelname)-8s  %(message)s",
        datefmt="%H:%M:%S",
    )

    ch = logging.StreamHandler()
    ch.setLevel(logging.DEBUG)
    ch.setFormatter(fmt)
    logger.addHandler(ch)

    log_path = Path(settings.data_dir) / "ingestion.log"
    try:
        log_path.parent.mkdir(parents=True, exist_ok=True)
        fh = logging.FileHandler(log_path, encoding="utf-8")
        fh.setLevel(logging.DEBUG)
        fh.setFormatter(fmt)
        logger.addHandler(fh)
    except OSError:
        pass

    return logger


# ---------------------------------------------------------------------------
# OCR — lazy-loaded via lru_cache (no global keyword)
# ---------------------------------------------------------------------------


@functools.lru_cache(maxsize=1)
def _get_paddle_ocr():
    from paddleocr import PaddleOCR

    log.info("Loading PaddleOCR (lang=%s) ...", settings.ocr_lang_paddle)
    return PaddleOCR(
        use_angle_cls=True,
        lang=settings.ocr_lang_paddle,
        show_log=False,
    )


@functools.lru_cache(maxsize=1)
def _get_surya_predictors():
    from surya.detection import DetectionPredictor
    from surya.recognition import RecognitionPredictor

    log.info("Loading Surya OCR ...")
    return (RecognitionPredictor(), DetectionPredictor())


def _ocr_image_paddle(image) -> str:
    ocr = _get_paddle_ocr()
    result = ocr.ocr(np.array(image), cls=True)
    lines = []
    for block in result or []:
        for entry in block or []:
            text = entry[1][0]
            if text and text.strip():
                lines.append(text.strip())
    return "\n".join(lines)


def _ocr_image_surya(image) -> str:
    rec_predictor, det_predictor = _get_surya_predictors()
    predictions = rec_predictor([image], [settings.ocr_lang_surya], det_predictor)
    lines = [line.text.strip() for line in predictions[0].text_lines if line.text.strip()]
    return "\n".join(lines)


def ocr_pdf_pages(doc, page_nums: list[int], filename: str) -> dict:
    results = {}
    zoom = settings.ocr_dpi / 72
    mat = fitz.Matrix(zoom, zoom)

    for page_num in page_nums:
        page = doc.load_page(page_num - 1)
        pix = page.get_pixmap(matrix=mat)
        img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)

        text = ""
        if settings.ocr_engine in ("paddleocr", "auto"):
            text = _ocr_image_paddle(img)

        if not text and settings.ocr_engine in ("surya", "auto"):
            text = _ocr_image_surya(img)

        results[page_num] = text

    return results


# ---------------------------------------------------------------------------
# PDF parsing
# ---------------------------------------------------------------------------


def _pymupdf_table_to_markdown(table) -> str:
    """Convert a PyMuPDF table to markdown format."""
    try:
        data = table.extract()
    except Exception:
        return ""
    if not data or len(data) < 1:
        return ""

    # Normalize column count
    max_cols = max(len(row) for row in data)
    rows = []
    for row in data:
        normalized = [str(cell).strip().replace("|", "\\|") if cell else "" for cell in row]
        while len(normalized) < max_cols:
            normalized.append("")
        rows.append(normalized)

    header = "| " + " | ".join(rows[0]) + " |"
    separator = "|" + "|".join(["---"] * max_cols) + "|"
    body = ["| " + " | ".join(row) + " |" for row in rows[1:]]
    return "\n".join([header, separator] + body) if body else header


def _extract_page_text(page) -> str:
    """Extract text from page blocks for better multi-column ordering."""
    blocks = page.get_text("blocks")
    blocks.sort(key=lambda b: (round(b[1] / 10) * 10, b[0]))
    return "\n".join(b[4] for b in blocks if len(b) > 4 and b[4].strip())


def _extract_page_tables(page, file_path: Path, page_num: int) -> list[Document]:
    """Detect and extract tables from a page. Returns list of table Documents."""
    docs: list[Document] = []
    try:
        tables = page.find_tables()
        if tables and tables.tables:
            for table in tables.tables:
                md_table = _pymupdf_table_to_markdown(table)
                if md_table:
                    docs.append(
                        Document(
                            page_content=md_table,
                            metadata={
                                "page": page_num,
                                "source": str(file_path),
                                "content_type": PageContentType.TABLE.value,
                            },
                        )
                    )
    except Exception:
        pass
    return docs


def _detect_tabular_heuristic(text: str, page_num: int) -> None:
    """Log warning if text has tabular-looking lines but find_tables() found nothing."""
    lines = text.split("\n")
    tabular_lines = sum(
        1 for line in lines if line.count("\t") >= 2 or (len(re.findall(r"  {3,}", line)) > 0)
    )
    if tabular_lines > 3:
        log.warning(
            "Page %d: %d tabular-looking lines but find_tables() found nothing",
            page_num,
            tabular_lines,
        )


def _should_ocr(text: str, min_chars: int, ocr_enabled: bool) -> bool:
    """Return True if OCR should be triggered based on text length and settings."""
    if not text and ocr_enabled:
        return True
    if text and len(text.strip()) < min_chars and ocr_enabled:
        return True
    return False


def _select_best_text(existing_text: str | None, ocr_text: str) -> str:
    """Choose between OCR text and existing short text layer, preferring longer."""
    if existing_text:
        if len(ocr_text) > len(existing_text) * 1.5:
            return ocr_text
        return clean_pdf_text(existing_text) or ocr_text
    return ocr_text


def _process_page_text(text: str, tables_found: bool, page_num: int, file_path: Path) -> Document | None:
    if text and not tables_found:
        _detect_tabular_heuristic(text, page_num)
    if not text or tables_found:
        return None
    cleaned = clean_pdf_text(text)
    if not cleaned:
        return None
    return Document(page_content=cleaned, metadata={"page": page_num, "source": str(file_path)})


def _process_ocr_result(
    page_num: int, ocr_text: str, text_to_compare: dict, file_path: Path
) -> Document | None:
    if not ocr_text:
        return None
    ocr_text = clean_pdf_text(ocr_text)
    if not ocr_text:
        return None
    final_text = _select_best_text(text_to_compare.get(page_num), ocr_text)
    return Document(page_content=final_text, metadata={"page": page_num, "source": str(file_path)})


def parse_pdf(file_path: Path) -> list[Document]:
    doc = fitz.open(str(file_path))
    pages = []

    ocr_pages_needed = []
    text_to_compare: dict[int, str] = {}

    for page_num in range(1, len(doc) + 1):
        page = doc.load_page(page_num - 1)
        text = _extract_page_text(page)

        table_docs = _extract_page_tables(page, file_path, page_num)
        pages.extend(table_docs)

        min_chars = settings.ocr_min_chars
        if _should_ocr(text, min_chars, settings.ocr_enabled):
            ocr_pages_needed.append(page_num)
            if text:
                text_to_compare[page_num] = text
        else:
            text_doc = _process_page_text(text, bool(table_docs), page_num, file_path)
            if text_doc:
                pages.append(text_doc)

    # --- Batch OCR ---
    if ocr_pages_needed:
        ocr_results = ocr_pdf_pages(doc, ocr_pages_needed, file_path.name)
        for page_num, ocr_text in ocr_results.items():
            ocr_doc = _process_ocr_result(page_num, ocr_text, text_to_compare, file_path)
            if ocr_doc:
                pages.append(ocr_doc)

    doc.close()
    return pages


# ---------------------------------------------------------------------------
# Text splitting
# ---------------------------------------------------------------------------


def merge_pdf_pages(pages: list[Document]) -> list[Document]:
    """Merge per-page Documents from the same source into a single Document.

    Preserves page range in metadata (page_start, page_end, pages list).
    Prevents the splitter from breaking text at page boundaries.
    """
    if len(pages) <= 1:
        return pages

    # Group by source
    by_source: dict[str, list[Document]] = {}
    for p in pages:
        src = p.metadata.get("source", "")
        by_source.setdefault(src, []).append(p)

    merged = []
    for _src, group in by_source.items():
        merged_text = "\n\n".join(p.page_content for p in group)
        has_pages = any("page" in p.metadata for p in group)
        meta = {**group[0].metadata}
        if has_pages:
            page_nums = sorted(p.metadata.get("page", i + 1) for i, p in enumerate(group))
            meta["page_start"] = page_nums[0]
            meta["page_end"] = page_nums[-1]
            meta["pages"] = page_nums
        merged.append(Document(page_content=merged_text, metadata=meta))
    return merged


def split_documents(docs: list[Document], domain: str = "general") -> list[Document]:
    """Split documents into chunks using structure-aware separators.

    Separators are chosen based on document domain:
    - legal: articles, sections, clauses (larger chunks)
    - general: markdown headers, paragraphs (standard chunks)

    Table chunks (content_type=table) are kept atomic — not split further.
    """
    if domain == DocDomain.LEGAL.value:
        return split_documents_legal(docs)

    separators = GENERAL_SEPARATORS

    # Separate table chunks (atomic) from text chunks
    tables = [d for d in docs if d.metadata.get("content_type") == PageContentType.TABLE.value]
    text_docs = [d for d in docs if d.metadata.get("content_type") != PageContentType.TABLE.value]

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.chunk_size,
        chunk_overlap=settings.chunk_overlap,
        length_function=len,
        separators=separators,
    )
    chunks = splitter.split_documents(text_docs)
    # Tables pass through unsplit
    chunks.extend(tables)
    log.info("Split %d documents into %d chunks (domain=%s)", len(docs), len(chunks), domain)
    return chunks


GENERAL_SEPARATORS = [
    "\n# ",
    "\n## ",
    "\n### ",
    "\n#### ",
    "\n\n",
    "\n",
    " ",
    "",
]


LEGAL_SEPARATORS = [
    "\nГлава ",
    "\nРаздел ",
    "\nЧасть ",
    "\nСтатья ",
    "\n§ ",
    "\nПункт ",
    "\n\n",
    "\n",
    " ",
    "",
]


def split_documents_legal(docs: list[Document]) -> list[Document]:
    """Split legal documents into chunks with larger size and legal-aware separators.

    Uses legal_chunk_size (default 1000) and legal_chunk_overlap (default 250)
    to keep articles/clauses intact.
    """
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.legal_chunk_size,
        chunk_overlap=settings.legal_chunk_overlap,
        length_function=len,
        separators=LEGAL_SEPARATORS,
    )
    chunks = splitter.split_documents(docs)
    log.info("Split %d legal documents into %d chunks", len(docs), len(chunks))
    return chunks


_ARTICLE_RE = re.compile(r"Статья\s+(\d+[\.\d]*)")


def extract_article_number(chunk_text: str) -> str | None:
    """Extract article number from chunk text if present."""
    m = _ARTICLE_RE.search(chunk_text)
    return m.group(1) if m else None


# ---------------------------------------------------------------------------
# Parsers for non-PDF formats
# ---------------------------------------------------------------------------


def _has_page_break(paragraph) -> bool:
    """Check if a paragraph contains a manual page break."""
    from docx.oxml.ns import qn

    for run in paragraph.runs:
        for br in run._element.findall(qn("w:br")):
            if br.get(qn("w:type")) == "page":
                return True
    return False


def _paragraph_list_prefix(paragraph) -> str | None:
    """Detect if paragraph is a list item and return indent prefix.

    Checks <w:numPr> in paragraph properties XML. Returns a prefix like
    "- " for level 0, "  - " for level 1, etc. Returns None if not a list item.
    """
    from docx.oxml.ns import qn

    pPr = paragraph._element.find(qn("w:pPr"))
    if pPr is None:
        return None
    numPr = pPr.find(qn("w:numPr"))
    if numPr is None:
        return None
    ilvl = numPr.find(qn("w:ilvl"))
    level = int(ilvl.get(qn("w:val"), "0")) if ilvl is not None else 0
    return "  " * level + "- "


def _extract_image_captions(doc) -> list[str]:
    """Extract alt text/descriptions from inline images in DOCX.

    Scans <w:drawing> elements for <wp:docPr descr="..."> attributes.
    """
    from docx.oxml.ns import qn

    captions = []
    for p in doc.paragraphs:
        for drawing in p._element.findall(f".//{qn('w:drawing')}"):
            for docPr in drawing.findall(f".//{qn('wp:docPr')}"):
                descr = docPr.get(qn("wp:descr")) or docPr.get("descr", "")
                if descr and descr.strip():
                    captions.append(f"[image: {descr.strip()}]")
    return captions


def _docx_table_to_markdown(table) -> str:
    """Convert a python-docx table to markdown table format.

    First row is used as header. Merged cells are handled by python-docx
    which returns the merged value for each cell in the range.
    """
    rows = []
    for row in table.rows:
        cells = [cell.text.strip().replace("|", "\\|") for cell in row.cells]
        rows.append(cells)
    if not rows:
        return ""

    # Normalize column count
    max_cols = max(len(r) for r in rows)
    for r in rows:
        while len(r) < max_cols:
            r.append("")

    header = "| " + " | ".join(rows[0]) + " |"
    separator = "|" + "|".join(["---"] * max_cols) + "|"
    body = ["| " + " | ".join(row) + " |" for row in rows[1:]]
    return "\n".join([header, separator] + body) if body else header


def _parse_docx(file_path: Path) -> tuple[str, dict]:
    """Parse DOCX and return (text, metadata).

    Detects page breaks via <w:br w:type="page"/> to provide page metadata.
    Tables are serialized as markdown tables.
    List items get restored prefixes (- ,   - ).
    Image alt text is extracted from inline drawings.
    """
    doc = docx.Document(str(file_path))
    parts = []
    page_numbers: list[int] = []
    current_page = 1

    for p in doc.paragraphs:
        if _has_page_break(p):
            current_page += 1
        if not p.text.strip():
            continue
        page_numbers.append(current_page)

        prefix = _paragraph_list_prefix(p)
        if prefix:
            parts.append(prefix + p.text)
        else:
            parts.append(p.text)

    for table in doc.tables:
        md_table = _docx_table_to_markdown(table)
        if md_table:
            parts.append("")  # blank line before table
            parts.append(md_table)

    # Extract image captions
    captions = _extract_image_captions(doc)
    for cap in captions:
        parts.append(cap)

    metadata: dict = {}
    if page_numbers:
        metadata["page_start"] = page_numbers[0]
        metadata["page_end"] = page_numbers[-1]
        metadata["pages"] = sorted(set(page_numbers))
    return "\n".join(parts), metadata


def _parse_rtf(file_path: Path) -> tuple[str, dict]:
    """Parse RTF and return (text, metadata)."""
    return rtf_to_text(file_path.read_text(encoding="utf-8", errors="replace")), {}


_MD_HEADER_RE = re.compile(r"^(#{1,6})\s+(.*)$", re.MULTILINE)
_MD_TABLE_RE = re.compile(r"^(\|.+\|)\s*$", re.MULTILINE)
_DATE_IN_FILENAME_RE = re.compile(r"(\d{4}-\d{2}-\d{2})")


def _clean_markdown_text(text: str) -> str:
    """Общая очистка markdown-разметки — используется и flat-парсером, и секционным."""
    text = re.sub(r"!\[([^\]]*)\]\(([^)]+)\)", r"[image: \1](\2)", text)
    text = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", text)
    text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)
    text = re.sub(r"[*_`~]+", "", text)
    text = html.unescape(text)
    return text.strip()


def _flush_text_segment(current_text: list[str], segments: list[tuple[str, str]]) -> None:
    if current_text:
        text = "\n".join(current_text).strip()
        if text:
            segments.append((PageContentType.TEXT.value, text))
        current_text.clear()


def _flush_table_segment(current_table: list[str], segments: list[tuple[str, str]]) -> None:
    if current_table:
        table_text = "\n".join(current_table).strip()
        if table_text:
            segments.append((PageContentType.TABLE.value, table_text))
        current_table.clear()


def _split_markdown_tables(content: str) -> list[tuple[str, str]]:
    """Split content into (content_type, text) segments.

    Detects markdown table blocks (lines starting and ending with |)
    and separates them from regular text. Returns list of (type, text)
    where type is 'text' or 'table'.
    """
    lines = content.split("\n")
    segments: list[tuple[str, str]] = []
    current_text: list[str] = []
    current_table: list[str] = []
    in_table = False

    for line in lines:
        is_table_line = bool(_MD_TABLE_RE.match(line.strip()))
        if is_table_line:
            if not in_table:
                _flush_text_segment(current_text, segments)
                in_table = True
            current_table.append(line)
        else:
            if in_table:
                _flush_table_segment(current_table, segments)
                in_table = False
            current_text.append(line)

    if in_table:
        _flush_table_segment(current_table, segments)
    else:
        _flush_text_segment(current_text, segments)

    return segments if segments else [(PageContentType.TEXT.value, content)]


def parse_markdown_sections(file_path: Path) -> list[tuple[str | None, str]]:
    r"""Разбивает markdown на (заголовок, контент) по структуре заголовков.

    Контент между двумя заголовками относится к предыдущему заголовку. Текст до
    первого заголовка (если есть) получает heading=None.

    Table blocks within sections are yielded as separate (heading, content) tuples
    with a special prefix '\x00TABLE:' to signal content_type: table downstream.
    """
    raw = file_path.read_text(encoding="utf-8", errors="replace")
    matches = list(_MD_HEADER_RE.finditer(raw))

    if not matches:
        segments = _split_markdown_tables(_clean_markdown_text(raw))
        sections: list[tuple[str | None, str]] = []
        for ctype, text in segments:
            prefix = "\x00TABLE:" if ctype == PageContentType.TABLE.value else ""
            sections.append((None, prefix + text))
        return sections

    result: list[tuple[str | None, str]] = []
    if matches[0].start() > 0:
        lead = _clean_markdown_text(raw[: matches[0].start()])
        if lead:
            for ctype, text in _split_markdown_tables(lead):
                prefix = "\x00TABLE:" if ctype == PageContentType.TABLE.value else ""
                result.append((None, prefix + text))

    for i, m in enumerate(matches):
        heading = m.group(2).strip()
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(raw)
        content = _clean_markdown_text(raw[start:end])
        if content:
            for ctype, text in _split_markdown_tables(content):
                prefix = "\x00TABLE:" if ctype == PageContentType.TABLE.value else ""
                result.append((heading, prefix + text))

    return result


def parse_docx_sections(file_path: Path) -> list[tuple[str | None, str]]:
    """Разбивает DOCX на (заголовок, контент) по параграфам со стилем Heading*/Title.

    Also detects page breaks to track page numbers within each section.
    """
    doc = docx.Document(str(file_path))
    sections: list[tuple[str | None, str]] = []
    current_heading: str | None = None
    current_lines: list[str] = []

    def _flush() -> None:
        content = "\n".join(current_lines).strip()
        if content:
            sections.append((current_heading, content))

    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        style_name = (p.style.name or "").lower() if p.style else ""
        if style_name.startswith("heading") or style_name == "title":
            _flush()
            current_heading = text
            current_lines = []
        else:
            current_lines.append(text)
    _flush()

    table_lines = []
    for table in doc.tables:
        md_table = _docx_table_to_markdown(table)
        if md_table:
            table_lines.append(md_table)
    if table_lines:
        sections.append((None, "\n\n".join(table_lines)))

    if not sections:
        text, _meta = _parse_docx(file_path)
        return [(None, text)]
    return sections


def extract_date_from_filename(filename: str) -> str | None:
    """Ищет дату вида YYYY-MM-DD в имени файла — простая эвристика для метаданных."""
    m = _DATE_IN_FILENAME_RE.search(filename)
    return m.group(1) if m else None


def _parse_markdown(file_path: Path) -> str:
    text = file_path.read_text(encoding="utf-8", errors="replace")
    return _clean_markdown_text(text)


def _parse_txt(file_path: Path) -> tuple[str, dict]:
    """Parse TXT and return (text, metadata)."""
    return file_path.read_text(encoding="utf-8", errors="replace"), {}


PARSERS = {
    ".docx": _parse_docx,
    ".doc": _parse_docx,
    ".rtf": _parse_rtf,
    ".md": _parse_markdown,
    ".txt": _parse_txt,
}


# ---------------------------------------------------------------------------
# Text cleaning
# ---------------------------------------------------------------------------


def clean_pdf_text(text: str) -> str:
    """Clean extracted PDF text: fix hyphenation, collapse whitespace, remove decorative lines."""
    # Fix hyphenation at line breaks
    text = re.sub(r"-\n", "", text)
    # Collapse multiple whitespace to single space
    text = re.sub(r"[^\S\n]+", " ", text)
    # Remove decorative separator lines (---, ===, *** etc.)
    text = re.sub(r"^[•\-=~*]{3,}\s*$", "", text, flags=re.MULTILINE)
    # Collapse multiple blank lines
    text = re.sub(r"\n{2,}", "\n", text)
    return text.strip()
