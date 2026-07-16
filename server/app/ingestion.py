"""
ingestion.py — индексирует документы в Qdrant.

Режимы запуска:
  python ingestion.py --docs_dir /app/docs           # добавить новые (не трогая уже загруженные)
  python ingestion.py --docs_dir /app/docs --reset   # сбросить и переиндексировать всё
  python ingestion.py --file /app/docs/report.pdf    # добавить один файл

Парсеры (нативные, без unstructured):
  .pdf  → PyMuPDF        .docx/.doc → python-docx + XML fallback
  .rtf  → striprtf       .md → markdown    .txt → plain read
"""

import argparse
import html
import json
import logging
import re
import time
from datetime import datetime
from pathlib import Path

from config import settings
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_qdrant import QdrantVectorStore
from qdrant_client import QdrantClient
from qdrant_client.models import Distance, VectorParams

# ---------------------------------------------------------------------------
# Логгер
# ---------------------------------------------------------------------------

def setup_logger() -> logging.Logger:
    logger = logging.getLogger("ingestion")
    if logger.handlers:
        return logger  # уже настроен (повторный вызов из FastAPI)

    logger.setLevel(logging.DEBUG)

    fmt = logging.Formatter(
        fmt="%(asctime)s  %(levelname)-8s  %(message)s",
        datefmt="%H:%M:%S",
    )

    # Консоль
    ch = logging.StreamHandler()
    ch.setLevel(logging.DEBUG)
    ch.setFormatter(fmt)
    logger.addHandler(ch)

    # Файл (рядом со скриптом)
    log_path = Path(__file__).parent / "ingestion.log"
    fh = logging.FileHandler(log_path, encoding="utf-8")
    fh.setLevel(logging.DEBUG)
    fh.setFormatter(fmt)
    logger.addHandler(fh)

    return logger


log = setup_logger()


# ---------------------------------------------------------------------------
# Реестр загруженных файлов
# Хранит: filename → {hash, chunks, indexed_at}
# Позволяет пропускать файлы которые уже есть в Qdrant
# ---------------------------------------------------------------------------

REGISTRY_PATH = Path(__file__).parent / "ingestion_registry.json"


def load_registry() -> dict:
    if REGISTRY_PATH.exists():
        try:
            return json.loads(REGISTRY_PATH.read_text(encoding="utf-8"))
        except Exception:
            return {}
    return {}


def save_registry(registry: dict):
    REGISTRY_PATH.write_text(
        json.dumps(registry, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


def file_hash(path: Path) -> str:
    """Быстрый хэш файла по размеру + дате изменения (не читаем содержимое)."""
    stat = path.stat()
    return f"{stat.st_size}_{int(stat.st_mtime)}"


def is_already_indexed(path: Path, registry: dict) -> bool:
    key = path.name
    if key not in registry:
        return False
    return registry[key].get("hash") == file_hash(path)


# ---------------------------------------------------------------------------
# PDF парсер
# ---------------------------------------------------------------------------

def parse_pdf(path: Path) -> str:
    try:
        import fitz
        doc = fitz.open(str(path))
        pages_text: dict[int, str] = {}
        scan_page_nums = []

        for i, page in enumerate(doc):
            blocks = page.get_text("blocks", sort=True)
            page_parts = [
                block[4].strip()
                for block in blocks
                if block[6] == 0 and block[4].strip()
            ]
            page_text = clean_pdf_text("\n".join(page_parts))

            if len(page_text.strip()) < 50:
                scan_page_nums.append(i)  # индекс страницы (0-based), OCR-им позже
            else:
                pages_text[i] = page_text

        if scan_page_nums:
            pages_str = ", ".join(str(p + 1) for p in scan_page_nums[:10])
            suffix = f"... (+{len(scan_page_nums) - 10})" if len(scan_page_nums) > 10 else ""
            log.warning("  Страницы без текстового слоя (скан?): %s%s", pages_str, suffix)

            if settings.ocr_enabled:
                ocr_results = ocr_pdf_pages(doc, scan_page_nums, path.name)
                pages_text.update(ocr_results)
            else:
                log.warning(
                    "  OCR выключен (OCR_ENABLED=false) — страницы-сканы пропущены: %s", path.name
                )

        doc.close()

        if not pages_text:
            raise RuntimeError(
                "PDF состоит из сканов, и OCR не смог извлечь текст (или отключён). "
                "Проверь качество скана / включи OCR_ENABLED=true."
            )

        ordered = [pages_text[i] for i in sorted(pages_text.keys())]
        return "\n\n".join(ordered)

    except RuntimeError:
        raise
    except Exception as e:
        raise RuntimeError(f"PDF parse error: {e}")


# ---------------------------------------------------------------------------
# OCR для страниц-сканов (PaddleOCR основной движок, Surya — опциональный)
# ---------------------------------------------------------------------------

_paddle_ocr_instance = None
_surya_predictors = None


def _get_paddle_ocr():
    global _paddle_ocr_instance
    if _paddle_ocr_instance is None:
        from paddleocr import PaddleOCR
        log.info("Загружаю PaddleOCR (lang=%s) ...", settings.ocr_lang_paddle)
        _paddle_ocr_instance = PaddleOCR(
            use_angle_cls=True,
            lang=settings.ocr_lang_paddle,
            show_log=False,
        )
    return _paddle_ocr_instance


def _get_surya_predictors():
    """Ленивая загрузка Surya. Импортируется только если явно запрошен движок surya/auto
    и пакет surya-ocr установлен (он не входит в базовый requirements.txt — см. README)."""
    global _surya_predictors
    if _surya_predictors is None:
        from surya.detection import DetectionPredictor
        from surya.recognition import RecognitionPredictor
        log.info("Загружаю Surya OCR ...")
        _surya_predictors = (RecognitionPredictor(), DetectionPredictor())
    return _surya_predictors


def _ocr_image_paddle(image) -> str:
    import numpy as np
    ocr = _get_paddle_ocr()
    result = ocr.ocr(np.array(image), cls=True)
    lines = []
    for block in result or []:
        for entry in block or []:
            text = entry[1][0]
            if text and text.strip():
                lines.append(text.strip())
    return "\n".join(lines)


def _ocr_image_surya(image) -> str:
    rec_predictor, det_predictor = _get_surya_predictors()
    predictions = rec_predictor([image], [settings.ocr_lang_surya], det_predictor)
    lines = [line.text.strip() for line in predictions[0].text_lines if line.text.strip()]
    return "\n".join(lines)


def ocr_pdf_pages(doc, page_nums: list[int], filename: str) -> dict:
    """
    Рендерит указанные страницы PDF в изображения и прогоняет через OCR.
    Возвращает {page_index: text}.
    """
    import fitz
    from PIL import Image

    results = {}
    zoom = settings.ocr_dpi / 72  # PyMuPDF рендерит в 72 DPI по умолчанию
    matrix = fitz.Matrix(zoom, zoom)

    for idx, page_num in enumerate(page_nums, 1):
        page = doc[page_num]
        pix = page.get_pixmap(matrix=matrix)
        image = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)

        log.info(
            "  OCR [%d/%d]  %s  стр. %d ...", idx, len(page_nums), filename, page_num + 1
        )

        text = ""
        engine_used = None
        try:
            if settings.ocr_engine in ("paddleocr", "auto"):
                text = _ocr_image_paddle(image)
                engine_used = "paddleocr"
            if (not text.strip()) and settings.ocr_engine in ("surya", "auto"):
                text = _ocr_image_surya(image)
                engine_used = "surya"
        except ImportError as e:
            log.error(
                "  OCR-движок '%s' не установлен (%s). "
                "См. requirements.txt / README для установки.",
                settings.ocr_engine, e,
            )
        except Exception as e:
            log.error("  Ошибка OCR на стр. %d: %s", page_num + 1, e)

        if text.strip():
            results[page_num] = clean_pdf_text(text)
            log.info(
                "  OCR [%d/%d]  OK (%s) — %d символов", idx, len(page_nums), engine_used, len(text)
            )
        else:
            log.warning("  OCR [%d/%d]  пусто — страница %d пропущена", idx, len(page_nums), page_num + 1)

    return results


def clean_pdf_text(text: str) -> str:
    import unicodedata
    text = unicodedata.normalize("NFKC", text)
    text = re.sub(r'(\w)-\n(\w)', r'\1\2', text)
    lines = text.split('\n')
    cleaned = []
    for line in lines:
        line = re.sub(r'[ \t]{2,}', ' ', line).strip()
        if line and not re.match(r'^[\s\-=_.|•·▪]+$', line):
            cleaned.append(line)
    return re.sub(r'\n{3,}', '\n\n', '\n'.join(cleaned)).strip()


# ---------------------------------------------------------------------------
# DOCX / DOC парсер
# ---------------------------------------------------------------------------

def parse_docx(path: Path) -> str:
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(str(path))
        parts = []
        for para in doc.paragraphs:
            if t := para.text.strip():
                parts.append(t)
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n\n".join(parts)
    except Exception as e:
        return _parse_docx_raw(path, original_error=str(e))


def _parse_docx_raw(path: Path, original_error: str = "") -> str:
    import xml.etree.ElementTree as ET
    import zipfile
    try:
        with zipfile.ZipFile(str(path), "r") as z:
            if "word/document.xml" not in z.namelist():
                raise RuntimeError("word/document.xml not found")
            with z.open("word/document.xml") as f:
                tree = ET.parse(f)
        ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        texts = [
            e.text for e in tree.getroot().iter(f"{{{ns}}}t")
            if e.text and e.text.strip()
        ]
        result = " ".join(texts)
        if not result.strip():
            raise RuntimeError("Пустой текст после XML-парсинга")
        return result
    except Exception as e:
        raise RuntimeError(f"DOCX parse failed (python-docx: {original_error}, raw XML: {e})")


# ---------------------------------------------------------------------------
# RTF / MD / TXT
# ---------------------------------------------------------------------------

def parse_rtf(path: Path) -> str:
    try:
        from striprtf.striprtf import rtf_to_text
        raw = path.read_bytes()
        for enc in ("utf-8", "cp1251", "cp1252", "latin-1"):
            try:
                text = raw.decode(enc); break
            except UnicodeDecodeError:
                continue
        else:
            text = raw.decode("utf-8", errors="replace")
        result = re.sub(r"\n{3,}", "\n\n", rtf_to_text(text)).strip()
        if not result:
            raise RuntimeError("Пустой текст после RTF-конвертации")
        return result
    except ImportError:
        raise RuntimeError("striprtf не установлен: pip install striprtf")
    except Exception as e:
        raise RuntimeError(f"RTF parse error: {e}")


def parse_md(path: Path) -> str:
    try:
        import markdown as md_lib
        raw = path.read_text(encoding="utf-8", errors="replace")
        html_content = md_lib.markdown(raw)
        text = re.sub(r"<[^>]+>", " ", html_content)
        return re.sub(r"\n{3,}", "\n\n", html.unescape(text)).strip()
    except ImportError:
        return path.read_text(encoding="utf-8", errors="replace")


def parse_txt(path: Path) -> str:
    for enc in ("utf-8", "cp1251", "cp1252", "latin-1"):
        try:
            return path.read_text(encoding=enc)
        except UnicodeDecodeError:
            continue
    return path.read_bytes().decode("utf-8", errors="replace")


# ---------------------------------------------------------------------------
# Диспетчер
# ---------------------------------------------------------------------------

PARSERS = {
    ".pdf":  parse_pdf,
    ".docx": parse_docx,
    ".doc":  parse_docx,
    ".rtf":  parse_rtf,
    ".md":   parse_md,
    ".txt":  parse_txt,
}


def parse_file(path: Path) -> Document | None:
    ext = path.suffix.lower()
    parser = PARSERS.get(ext)
    if parser is None:
        log.debug("  SKIP  неподдерживаемый формат: %s", path.name)
        return None
    try:
        text = parser(path)
        if not text or len(text.strip()) < 20:
            log.warning("  SKIP  слишком мало текста: %s", path.name)
            return None
        return Document(
            page_content=text,
            metadata={
                "source":     str(path),
                "filename":   path.name,
                "extension":  ext,
                "size_bytes": path.stat().st_size,
            },
        )
    except Exception as e:
        log.error("  ERROR %s: %s", path.name, e)
        return None


# ---------------------------------------------------------------------------
# Загрузка документов из папки
# ---------------------------------------------------------------------------

def load_documents(
    docs_dir: str,
    registry: dict,
    force: bool = False,
) -> tuple[list[Document], int]:
    """
    Возвращает (список документов, количество пропущенных).
    force=True — игнорирует реестр, перезагружает всё.
    """
    docs_path = Path(docs_dir)
    if not docs_path.exists():
        log.error("Папка не найдена: %s", docs_dir)
        return [], 0

    all_files = sorted(
        f for f in docs_path.rglob("*")
        if f.is_file() and f.suffix.lower() in PARSERS
    )

    log.info("Найдено файлов: %d в %s", len(all_files), docs_dir)

    documents, skipped_cached, ok, errors = [], 0, 0, 0

    for i, file_path in enumerate(all_files, 1):
        prefix = f"[{i:>3}/{len(all_files)}]"

        if not force and is_already_indexed(file_path, registry):
            log.info("%s CACHED  %s", prefix, file_path.name)
            skipped_cached += 1
            continue

        size_kb = file_path.stat().st_size / 1024
        log.info("%s PARSE   %s  (%.1f KB)", prefix, file_path.name, size_kb)
        t0 = time.monotonic()

        doc = parse_file(file_path)
        elapsed = time.monotonic() - t0

        if doc:
            documents.append(doc)
            log.info(
                "%s OK      %s — %s символов, %d страниц approx, %.2fs",
                prefix, file_path.name,
                f"{len(doc.page_content):,}",
                doc.page_content.count("\n\n") + 1,
                elapsed,
            )
            ok += 1
        else:
            errors += 1

    log.info(
        "Парсинг завершён: ✓ %d загружено, ✗ %d ошибок, ○ %d уже в реестре",
        ok, errors, skipped_cached,
    )
    return documents, skipped_cached


# ---------------------------------------------------------------------------
# Сплиттер
# ---------------------------------------------------------------------------

def split_documents(docs: list[Document]) -> list[Document]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.chunk_size,
        chunk_overlap=settings.chunk_overlap,
        separators=["\n\n", "\n", ". ", ", ", " ", ""],
        length_function=len,
    )
    chunks = splitter.split_documents(docs)
    log.info("Сплиттер: %d чанков из %d документов", len(chunks), len(docs))
    return chunks


# ---------------------------------------------------------------------------
# Qdrant helpers
# ---------------------------------------------------------------------------

def get_embeddings() -> HuggingFaceEmbeddings:
    log.info("Загружаю эмбеддинг-модель %s ...", settings.embed_model)
    emb = HuggingFaceEmbeddings(
        model_name=settings.embed_model,
        model_kwargs={"device": "cpu"},
        encode_kwargs={"normalize_embeddings": True},
    )
    log.info("Модель загружена")
    return emb


def ensure_collection(client: QdrantClient, vector_size: int, reset: bool = False):
    existing = [c.name for c in client.get_collections().collections]
    if settings.collection_name in existing:
        if reset:
            log.info("Удаляю коллекцию '%s' ...", settings.collection_name)
            client.delete_collection(settings.collection_name)
        else:
            info = client.get_collection(settings.collection_name)
            count = info.points_count or 0
            log.info(
                "Коллекция '%s' существует — %d точек. Добавляю новые документы.",
                settings.collection_name, count,
            )
            return
    log.info("Создаю коллекцию '%s' (dim=%d) ...", settings.collection_name, vector_size)
    client.create_collection(
        collection_name=settings.collection_name,
        vectors_config=VectorParams(size=vector_size, distance=Distance.COSINE),
    )


def upload_to_qdrant(chunks: list[Document], embeddings: HuggingFaceEmbeddings):
    batch_size = 100
    total = len(chunks)
    log.info("Загружаю %d чанков в Qdrant батчами по %d ...", total, batch_size)
    t0 = time.monotonic()

    for i in range(0, total, batch_size):
        batch = chunks[i:i + batch_size]
        QdrantVectorStore.from_documents(
            documents=batch,
            embedding=embeddings,
            url=settings.qdrant_url,
            collection_name=settings.collection_name,
            force_recreate=False,
        )
        done = min(i + batch_size, total)
        elapsed = time.monotonic() - t0
        speed = done / elapsed if elapsed > 0 else 0
        eta = (total - done) / speed if speed > 0 else 0
        log.info(
            "  Загружено %d/%d чанков  (%.1f ч/с, ETA ~%.0fs)",
            done, total, speed, eta,
        )

    log.info("Загрузка в Qdrant завершена за %.1fs", time.monotonic() - t0)


# ---------------------------------------------------------------------------
# Публичные функции — вызываются из FastAPI и CLI
# ---------------------------------------------------------------------------

def run_ingestion(docs_dir: str, reset: bool = False):
    """Полная индексация папки. reset=True — пересоздаёт коллекцию."""
    t_start = time.monotonic()
    log.info("=" * 55)
    log.info("RAG Ingestion  |  режим: %s", "RESET" if reset else "APPEND")
    log.info("docs_dir : %s", docs_dir)
    log.info("model    : %s", settings.embed_model)
    log.info("qdrant   : %s  /  collection: %s", settings.qdrant_url, settings.collection_name)
    log.info("=" * 55)

    registry = load_registry()
    if reset:
        registry = {}

    embeddings = get_embeddings()
    vector_size = len(embeddings.embed_query("тест"))
    qdrant_client = QdrantClient(url=settings.qdrant_url)
    ensure_collection(qdrant_client, vector_size, reset=reset)

    docs, cached = load_documents(docs_dir, registry, force=reset)
    if not docs:
        if cached > 0:
            log.info("Все файлы уже в реестре — нечего индексировать. Передай --reset чтобы переиндексировать.")
        else:
            log.error("Ни одного документа не загружено. Проверь папку и форматы.")
        return

    chunks = split_documents(docs)
    upload_to_qdrant(chunks, embeddings)

    # Обновляем реестр
    for doc in docs:
        path = Path(doc.metadata["source"])
        registry[path.name] = {
            "hash":       file_hash(path),
            "source":     str(path),
            "chunks":     sum(1 for c in chunks if c.metadata.get("source") == str(path)),
            "chars":      len(doc.page_content),
            "indexed_at": datetime.now().isoformat(timespec="seconds"),
        }
    save_registry(registry)

    total_elapsed = time.monotonic() - t_start
    log.info("=" * 55)
    log.info("✓ ГОТОВО  |  %d чанков  |  %.1fs всего", len(chunks), total_elapsed)
    log.info("Реестр: %d файлов в %s", len(registry), REGISTRY_PATH)
    log.info("=" * 55)


def run_ingest_file(file_path: str):
    """Добавить один файл в существующую коллекцию (не сбрасывая остальное)."""
    t_start = time.monotonic()
    path = Path(file_path)

    log.info("=" * 55)
    log.info("RAG Ingestion  |  режим: SINGLE FILE")
    log.info("file     : %s", path)
    log.info("=" * 55)

    if not path.exists():
        log.error("Файл не найден: %s", file_path)
        return
    if path.suffix.lower() not in PARSERS:
        log.error("Неподдерживаемый формат: %s", path.suffix)
        return

    registry = load_registry()

    # Проверяем не загружен ли уже этот файл (по хэшу)
    if is_already_indexed(path, registry):
        log.warning(
            "Файл '%s' уже в реестре с тем же хэшем. "
            "Передай --reset-file чтобы переиндексировать.", path.name
        )
        return

    log.info("PARSE   %s  (%.1f KB)", path.name, path.stat().st_size / 1024)
    doc = parse_file(path)
    if not doc:
        log.error("Не удалось распарсить файл.")
        return

    log.info("OK  %s  —  %s символов", path.name, f"{len(doc.page_content):,}")

    embeddings = get_embeddings()
    vector_size = len(embeddings.embed_query("тест"))
    qdrant_client = QdrantClient(url=settings.qdrant_url)
    ensure_collection(qdrant_client, vector_size, reset=False)

    chunks = split_documents([doc])
    upload_to_qdrant(chunks, embeddings)

    registry[path.name] = {
        "hash":       file_hash(path),
        "source":     str(path),
        "chunks":     len(chunks),
        "chars":      len(doc.page_content),
        "indexed_at": datetime.now().isoformat(timespec="seconds"),
    }
    save_registry(registry)

    log.info("=" * 55)
    log.info("✓ ГОТОВО  |  %d чанков  |  %.1fs", len(chunks), time.monotonic() - t_start)
    log.info("=" * 55)


def list_registry():
    """Показать список проиндексированных файлов."""
    registry = load_registry()
    if not registry:
        log.info("Реестр пуст — ни одного файла не проиндексировано.")
        return
    log.info("Проиндексированных файлов: %d", len(registry))
    log.info("%-50s  %6s  %8s  %s", "Файл", "Чанков", "Символов", "Дата")
    log.info("-" * 85)
    for name, meta in sorted(registry.items()):
        log.info(
            "%-50s  %6s  %8s  %s",
            name[:50],
            meta.get("chunks", "?"),
            f"{meta.get('chars', 0):,}",
            meta.get("indexed_at", "?")[:19],
        )


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="RAG Ingestion Pipeline",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Примеры:
  # Добавить только новые файлы (уже загруженные пропускаются)
  python ingestion.py --docs_dir /code/project/docs

  # Переиндексировать всё с нуля
  python ingestion.py --docs_dir /code/project/docs --reset

  # Добавить один файл
  python ingestion.py --file /code/project/docs/новый_документ.pdf

  # Показать что уже проиндексировано
  python ingestion.py --list
        """,
    )
    parser.add_argument("--docs_dir", help="Папка с документами")
    parser.add_argument("--file",     help="Добавить один конкретный файл")
    parser.add_argument("--reset",    action="store_true", help="Сбросить коллекцию и реестр, переиндексировать всё")
    parser.add_argument("--reset-file", action="store_true", help="Переиндексировать файл даже если он уже в реестре")
    parser.add_argument("--list",     action="store_true", help="Показать список проиндексированных файлов")
    args = parser.parse_args()

    if args.list:
        list_registry()
    elif args.file:
        if args.reset_file:
            # Удалим из реестра чтобы force-переиндексировать
            registry = load_registry()
            registry.pop(Path(args.file).name, None)
            save_registry(registry)
        run_ingest_file(args.file)
    elif args.docs_dir:
        run_ingestion(args.docs_dir, reset=args.reset)
    else:
        parser.print_help()
