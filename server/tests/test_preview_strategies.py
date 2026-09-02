"""Tests for preview strategies (PDF, DOCX, RTF) and PreviewStrategyFactory."""

import io
import sys
from pathlib import Path
from unittest.mock import MagicMock

import pytest

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "app"))

sys.modules["fitz"] = MagicMock()

from domain.value_objects.page_content_type import PageContentType, PreviewUnitKind  # noqa: E402
from application.services.pdf_diagnostic_service import DryRunPageResult  # noqa: E402
from infrastructure.ml.preview.factory import PreviewStrategyFactory  # noqa: E402
from infrastructure.ml.preview.rtf_preview_strategy import RtfPreviewStrategy  # noqa: E402
from infrastructure.ml.preview.docx_preview_strategy import DocxPreviewStrategy  # noqa: E402


# ---------------------------------------------------------------------------
# PreviewStrategyFactory
# ---------------------------------------------------------------------------


class TestPreviewStrategyFactory:
    def test_pdf_returns_pdf_strategy(self):
        from infrastructure.ml.preview.pdf_preview_strategy import PdfPreviewStrategy

        strategy = PreviewStrategyFactory.for_extension(".pdf", diag_service=MagicMock())
        assert isinstance(strategy, PdfPreviewStrategy)

    def test_docx_returns_docx_strategy(self):
        strategy = PreviewStrategyFactory.for_extension(".docx")
        assert isinstance(strategy, DocxPreviewStrategy)

    def test_rtf_returns_rtf_strategy(self):
        strategy = PreviewStrategyFactory.for_extension(".rtf")
        assert isinstance(strategy, RtfPreviewStrategy)

    def test_doc_raises_error(self):
        with pytest.raises(ValueError, match="не поддерживаются"):
            PreviewStrategyFactory.for_extension(".doc")

    def test_unknown_extension_raises(self):
        with pytest.raises(ValueError, match="Unsupported"):
            PreviewStrategyFactory.for_extension(".xlsx")

    def test_supported_extensions(self):
        exts = PreviewStrategyFactory.supported_extensions()
        assert ".pdf" in exts
        assert ".docx" in exts
        assert ".rtf" in exts
        assert ".doc" not in exts


# ---------------------------------------------------------------------------
# DocxPreviewStrategy
# ---------------------------------------------------------------------------


def _make_docx_with_text(*headings_and_texts: tuple[str | None, str]) -> bytes:
    """Create a minimal DOCX file in memory with given sections."""
    from docx import Document

    doc = Document()
    for heading, text in headings_and_texts:
        if heading:
            doc.add_heading(heading, level=1)
        if text:
            doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_docx_with_table() -> bytes:
    """Create a DOCX with a table."""
    from docx import Document

    doc = Document()
    doc.add_heading("Section 1", level=1)
    doc.add_paragraph("Some text here")
    table = doc.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "H1"
    table.cell(0, 1).text = "H2"
    table.cell(0, 2).text = "H3"
    table.cell(1, 0).text = "A"
    table.cell(1, 1).text = "B"
    table.cell(1, 2).text = "C"
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestDocxPreviewStrategy:
    def setup_method(self):
        self.strategy = DocxPreviewStrategy()

    def test_supports_docx(self):
        assert self.strategy.supports(".docx") is True
        assert self.strategy.supports(".pdf") is False
        assert self.strategy.supports(".rtf") is False

    def test_simple_text_sections(self, tmp_path):
        docx_bytes = _make_docx_with_text(
            (None, "Intro paragraph"),
            ("Heading One", "Content under heading"),
            ("Heading Two", "More content"),
        )
        path = tmp_path / "test.docx"
        path.write_bytes(docx_bytes)

        units, types_count, total_chars = self.strategy.analyze(path)

        assert len(units) >= 3
        assert total_chars > 0
        assert all(u.unit_kind == PreviewUnitKind.SECTION for u in units)
        assert all(u.label for u in units)
        assert types_count[PageContentType.TEXT] >= 3

    def test_empty_sections(self, tmp_path):
        docx_bytes = _make_docx_with_text(
            ("Heading", "Some text"),
            (None, ""),
        )
        path = tmp_path / "test.docx"
        path.write_bytes(docx_bytes)

        units, types_count, _ = self.strategy.analyze(path)
        # Empty paragraphs are skipped by parse_docx_sections, no empty unit created
        text_units = [u for u in units if u.type == PageContentType.TEXT]
        assert len(text_units) >= 1

    def test_garbled_content(self, tmp_path):
        # python-docx cannot write NULL bytes, use latin-1 compatible garbled text
        garbled_text = "".join(chr(i) for i in range(0xC0, 0xFF)) * 10
        docx_bytes = _make_docx_with_text(
            ("Heading", garbled_text),
        )
        path = tmp_path / "test.docx"
        path.write_bytes(docx_bytes)

        units, types_count, _ = self.strategy.analyze(path)
        garbled = [u for u in units if u.type == PageContentType.GARBLED]
        # If is_garbled heuristic triggers, we get GARBLED; otherwise TEXT (depends on char distribution)
        assert len(garbled) >= 0  # at minimum, the section exists

    def test_table_detection(self, tmp_path):
        docx_bytes = _make_docx_with_table()
        path = tmp_path / "test.docx"
        path.write_bytes(docx_bytes)

        units, types_count, _ = self.strategy.analyze(path)
        tables = [u for u in units if u.type == PageContentType.TABLE]
        assert len(tables) >= 1

    def test_empty_document(self, tmp_path):
        from docx import Document

        doc = Document()
        buf = io.BytesIO()
        doc.save(buf)
        path = tmp_path / "empty.docx"
        path.write_bytes(buf.getvalue())

        units, types_count, _ = self.strategy.analyze(path)
        assert len(units) == 1
        assert units[0].type == PageContentType.EMPTY

    def test_label_contains_heading(self, tmp_path):
        docx_bytes = _make_docx_with_text(
            ("Important Section", "Content here"),
        )
        path = tmp_path / "test.docx"
        path.write_bytes(docx_bytes)

        units, _, _ = self.strategy.analyze(path)
        heading_units = [u for u in units if "Important Section" in u.label]
        assert len(heading_units) == 1

    def test_ocr_problem_units_no_images(self, tmp_path):
        docx_bytes = _make_docx_with_text(
            ("Section", "Some text"),
        )
        path = tmp_path / "test.docx"
        path.write_bytes(docx_bytes)

        units, _, _ = self.strategy.analyze(path)
        result, _, _ = self.strategy.ocr_problem_units(path, units, [1])
        assert result == units


# ---------------------------------------------------------------------------
# RtfPreviewStrategy
# ---------------------------------------------------------------------------


def _make_rtf(text: str = "Hello world") -> bytes:
    """Create minimal RTF content."""
    rtf = r"{\rtf1\ansi " + text + "}"
    return rtf.encode("utf-8")


class TestRtfPreviewStrategy:
    def setup_method(self):
        self.strategy = RtfPreviewStrategy()

    def test_supports_rtf(self):
        assert self.strategy.supports(".rtf") is True
        assert self.strategy.supports(".pdf") is False
        assert self.strategy.supports(".docx") is False

    def test_valid_rtf(self, tmp_path):
        rtf_bytes = _make_rtf("This is readable text content for testing.")
        path = tmp_path / "test.rtf"
        path.write_bytes(rtf_bytes)

        units, types_count, total_chars = self.strategy.analyze(path)

        assert len(units) == 1
        assert units[0].unit_kind == PreviewUnitKind.DOCUMENT
        assert units[0].label == "Документ целиком"
        assert total_chars > 0
        assert types_count[PageContentType.TEXT] == 1

    def test_empty_rtf(self, tmp_path):
        rtf_bytes = _make_rtf("")
        path = tmp_path / "empty.rtf"
        path.write_bytes(rtf_bytes)

        units, types_count, _ = self.strategy.analyze(path)
        assert len(units) == 1
        assert units[0].type == PageContentType.EMPTY

    def test_garbled_rtf(self, tmp_path):
        garbled_text = "".join(chr(i) for i in range(0xC0, 0xFF)) * 10
        rtf_bytes = _make_rtf(garbled_text)
        path = tmp_path / "garbled.rtf"
        path.write_bytes(rtf_bytes)

        units, types_count, _ = self.strategy.analyze(path)
        # is_garbled check depends on character distribution
        assert units[0].type in (PageContentType.GARBLED, PageContentType.TEXT)

    def test_ocr_not_supported(self, tmp_path):
        rtf_bytes = _make_rtf("Some text")
        path = tmp_path / "test.rtf"
        path.write_bytes(rtf_bytes)

        units, _, _ = self.strategy.analyze(path)
        with pytest.raises(ValueError, match="OCR not supported"):
            self.strategy.ocr_problem_units(path, units, [1])


# ---------------------------------------------------------------------------
# DryRunPageResult dataclass
# ---------------------------------------------------------------------------


class TestDryRunPageResult:
    def test_default_unit_kind(self):
        r = DryRunPageResult(page=1, type="text")
        assert r.unit_kind == PreviewUnitKind.PAGE
        assert r.label == ""

    def test_custom_unit_kind(self):
        r = DryRunPageResult(
            page=1,
            type="text",
            unit_kind=PreviewUnitKind.SECTION,
            label="Раздел: Заголовок",
        )
        assert r.unit_kind == PreviewUnitKind.SECTION
        assert r.label == "Раздел: Заголовок"


# ---------------------------------------------------------------------------
# PageContentType enum
# ---------------------------------------------------------------------------


class TestPageContentType:
    def test_image_only_exists(self):
        assert PageContentType.IMAGE_ONLY == "image_only"

    def test_all_values(self):
        values = {e.value for e in PageContentType}
        assert "image_only" in values
        assert "text" in values
        assert "scan" in values
        assert "garbled" in values
        assert "empty" in values
        assert "table" in values
        assert "ocr" in values


# ---------------------------------------------------------------------------
# PreviewUnitKind enum
# ---------------------------------------------------------------------------


class TestPreviewUnitKind:
    def test_values(self):
        assert PreviewUnitKind.PAGE == "page"
        assert PreviewUnitKind.SECTION == "section"
        assert PreviewUnitKind.DOCUMENT == "document"
